"""Word document exporter — generates .docx reports from analysis results."""

import re
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def _add_styles(doc: Document):
    """Add custom styles to the document."""
    # Title style
    style = doc.styles["Title"]
    font = style.font
    font.size = Pt(24)
    font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # DocAI blue

    # Subtitle
    style = doc.styles["Subtitle"]
    font = style.font
    font.size = Pt(14)
    font.color.rgb = RGBColor(0x58, 0x58, 0x58)


def _parse_markdown_to_docx(doc: Document, markdown_text: str):
    """Convert markdown-formatted text into Word document elements."""
    lines = markdown_text.split("\n")
    in_table = False
    table_rows = []
    in_code_block = False

    for line in lines:
        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            p = doc.add_paragraph(stripped)
            p.style = doc.styles["No Spacing"]
            run = p.runs[0] if p.runs else p.add_run(stripped)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        # Table detection
        if "|" in stripped and stripped.startswith("|"):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            # Skip separator rows
            if all(c.replace("-", "").replace(":", "") == "" for c in cells):
                continue
            table_rows.append(cells)
            in_table = True
            continue
        elif in_table:
            # End of table — render it
            _render_table(doc, table_rows)
            table_rows = []
            in_table = False

        # Empty line
        if not stripped:
            doc.add_paragraph("")
            continue

        # Headings
        if stripped.startswith("##### "):
            doc.add_heading(stripped[6:], level=5)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, text)
        # Numbered lists
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, text)
        # Bold line (like **Key Points**)
        elif stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
        # Italic metadata line
        elif stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.italic = True
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        # Regular paragraph
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, stripped)

    # Flush remaining table
    if table_rows:
        _render_table(doc, table_rows)


def _add_formatted_text(paragraph, text: str):
    """Add text with basic bold/italic formatting."""
    # Split on **bold** and *italic* patterns
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def _render_table(doc: Document, rows: list[list[str]]):
    """Render a list of rows as a Word table."""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Light Grid Accent 1"

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = cell_text
                # Bold header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    doc.add_paragraph("")  # spacing after table


def export_to_word(
    analysis_result: str,
    folder_path: str,
    instruction: str,
    chunk_result=None,
    output_path: str | None = None,
) -> str:
    """Export analysis result to a Word document.

    Args:
        analysis_result: The LLM's markdown response.
        folder_path: Source folder that was analyzed.
        instruction: The user's instruction.
        chunk_result: Optional ChunkResult for metadata.
        output_path: Where to save. Auto-generates if None.

    Returns:
        Path to the saved .docx file.
    """
    doc = Document()
    _add_styles(doc)

    # Title page
    doc.add_heading("DocAI Analysis Report", level=0)

    # Metadata
    p = doc.add_paragraph()
    p.add_run("Generated: ").bold = True
    p.add_run(datetime.now().strftime("%B %d, %Y at %I:%M %p"))

    p = doc.add_paragraph()
    p.add_run("Source Folder: ").bold = True
    p.add_run(folder_path)

    p = doc.add_paragraph()
    p.add_run("Instruction: ").bold = True
    p.add_run(instruction)

    if chunk_result:
        p = doc.add_paragraph()
        p.add_run("Documents Analyzed: ").bold = True
        p.add_run(str(len(chunk_result.chunks)))

        p = doc.add_paragraph()
        p.add_run("Total Tokens: ").bold = True
        p.add_run(f"{chunk_result.total_tokens:,}")

        strategy = "Single-shot" if chunk_result.fits_single_shot else "Map-reduce"
        p = doc.add_paragraph()
        p.add_run("Strategy: ").bold = True
        p.add_run(strategy)

        # Document list
        doc.add_heading("Documents Processed", level=2)
        for c in chunk_result.chunks:
            doc.add_paragraph(
                f"{c.source} ({c.tokens:,} tokens)", style="List Bullet"
            )

    # Divider
    doc.add_paragraph("")
    doc.add_heading("Analysis Results", level=1)

    # Parse and add the analysis content
    _parse_markdown_to_docx(doc, analysis_result)

    # Footer
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generated by DocAI")
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # Save
    if not output_path:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"DocAI_Report_{timestamp}.docx"

    doc.save(output_path)
    return output_path
