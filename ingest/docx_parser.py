"""Word document parser — extracts text, tables, and metadata from .docx files."""

from dataclasses import dataclass, field
from pathlib import Path
from docx import Document


@dataclass
class ParsedDocument:
    filename: str
    text: str
    tables: list[list[list[str]]] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)


def parse_docx(path: Path) -> ParsedDocument:
    """Extract text, tables, and metadata from a Word document."""
    doc = Document(str(path))

    # Extract paragraphs preserving structure
    paragraphs = []
    for p in doc.paragraphs:
        if p.text.strip():
            # Preserve heading structure
            if p.style and p.style.name.startswith("Heading"):
                level = p.style.name.replace("Heading ", "")
                paragraphs.append(f"{'#' * int(level)} {p.text}")
            else:
                paragraphs.append(p.text)

    text = "\n\n".join(paragraphs)

    # Extract tables
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append(rows)

    # Extract metadata
    props = doc.core_properties
    metadata = {
        "title": props.title or "",
        "author": props.author or "",
        "created": str(props.created) if props.created else "",
        "modified": str(props.modified) if props.modified else "",
    }

    return ParsedDocument(
        filename=path.name,
        text=text,
        tables=tables,
        metadata=metadata,
    )
