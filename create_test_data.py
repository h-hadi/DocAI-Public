#!/usr/bin/env python3
"""Generate test data for DocAI pipeline verification."""

from pathlib import Path
import random

# Ensure test directories exist
test_dir = Path(__file__).parent / "test_data"
yoy_dir = test_dir / "financials"
summary_dir = test_dir / "reports"
yoy_dir.mkdir(parents=True, exist_ok=True)
summary_dir.mkdir(parents=True, exist_ok=True)

# ============================================================
# 1. Year-over-Year Excel test data (5 files, one per year)
# ============================================================
try:
    import openpyxl

    for year in range(2020, 2025):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Financial Summary"

        # Headers
        headers = ["Category", "Q1", "Q2", "Q3", "Q4", "Annual Total"]
        for col, h in enumerate(headers, 1):
            ws.cell(row=1, column=col, value=h)

        # Data with realistic growth + one anomaly in 2023
        base_revenue = 1_000_000 + (year - 2020) * 150_000
        base_expense = 600_000 + (year - 2020) * 80_000

        categories = {
            "Revenue": base_revenue,
            "Operating Expenses": base_expense,
            "Salaries": base_expense * 0.6,
            "Travel": 50_000 + (year - 2020) * 5_000,
            "Equipment": 30_000 + (year - 2020) * 3_000,
            "Grants Received": 200_000 + (year - 2020) * 25_000,
            "Donations": 80_000 + (year - 2020) * 10_000,
        }

        # Inject anomaly: 2023 Travel spikes 300%
        if year == 2023:
            categories["Travel"] = 200_000  # 3x normal

        for row, (cat, annual) in enumerate(categories.items(), 2):
            ws.cell(row=row, column=1, value=cat)
            quarterly = []
            for q in range(4):
                # Add some variance per quarter
                q_val = round(annual / 4 * (0.9 + random.random() * 0.2), 2)
                quarterly.append(q_val)
                ws.cell(row=row, column=q + 2, value=q_val)
            ws.cell(row=row, column=6, value=round(sum(quarterly), 2))

        wb.save(yoy_dir / f"financials_{year}.xlsx")

    print(f"Created 5 Excel files in {yoy_dir}")

except ImportError:
    print("openpyxl not installed — skipping Excel test data")


# ============================================================
# 2. Word document test data for Summarizer
# ============================================================
try:
    from docx import Document

    # Report 1: Annual Report
    doc = Document()
    doc.add_heading("Research Institute Annual Report 2024", level=1)
    doc.add_paragraph(
        "The Research Institute continued its mission of "
        "advancing ideas promoting economic prosperity, national security, and democratic "
        "governance. Key achievements this year include the digitization of 50,000 "
        "archival documents from the MENA collection, expansion of the oral history "
        "program to include 30 new interviews, and the launch of the Digital Archivist "
        "Initiative."
    )
    doc.add_heading("Collection Highlights", level=2)
    doc.add_paragraph(
        "The MENA collection acquired 15 new collections totaling 200 linear feet "
        "of material. Notable acquisitions include personal papers of diplomatic "
        "figures from the Gulf region and photographic archives documenting social "
        "movements across North Africa."
    )
    doc.add_heading("Digital Initiatives", level=2)
    doc.add_paragraph(
        "The AI-powered discovery tools developed in collaboration with the library "
        "team processed over 2 million metadata records. OCR accuracy for Arabic-script "
        "documents improved from 72% to 89% through custom training datasets."
    )
    doc.add_heading("Financial Overview", level=2)
    doc.add_paragraph(
        "Total operating budget: $2.4M. Revenue from grants: $1.2M. "
        "Donor contributions: $800K. Operating surplus: $200K."
    )
    doc.save(summary_dir / "annual_report_2024.docx")

    # Report 2: Strategic Plan
    doc = Document()
    doc.add_heading("Digital Preservation Strategic Plan 2025-2027", level=1)
    doc.add_paragraph(
        "This strategic plan outlines the Research Institute's approach to "
        "digital preservation over the next three years. The plan addresses "
        "three key areas: infrastructure modernization, AI-enhanced cataloging, "
        "and community access expansion."
    )
    doc.add_heading("Infrastructure", level=2)
    doc.add_paragraph(
        "Migration from on-premises storage to hybrid cloud architecture. "
        "Target: 99.99% availability for digital collections. Investment: $500K "
        "over 3 years for AWS infrastructure and redundancy."
    )
    doc.add_heading("AI-Enhanced Cataloging", level=2)
    doc.add_paragraph(
        "Deploy machine learning models for automatic metadata extraction from "
        "scanned documents. Target: reduce cataloging time by 60%. Technologies: "
        "OCR, named entity recognition, subject classification."
    )
    doc.add_heading("Community Access", level=2)
    doc.add_paragraph(
        "Launch public portal for MENA collections. Target: 100,000 unique visitors "
        "in first year. Include multi-language search (Arabic, Farsi, Turkish, English)."
    )
    doc.save(summary_dir / "strategic_plan_2025.docx")

    # Report 3: Collection Assessment
    doc = Document()
    doc.add_heading("MENA Collection Assessment Report", level=1)
    doc.add_paragraph(
        "Assessment of the Middle East and North Africa archival holdings "
        "at the Research Institute, covering extent, condition, access level, "
        "and digitization priority."
    )
    doc.add_heading("Extent", level=2)
    doc.add_paragraph(
        "Total holdings: 450 collections, 3,200 linear feet. "
        "Digitized: 15% (480 linear feet). "
        "Priority for digitization: 85 collections identified as high-value."
    )
    doc.add_heading("Condition", level=2)
    doc.add_paragraph(
        "Good: 70%. Fair: 20%. Poor: 10%. "
        "Items in poor condition include early 20th century photographs and "
        "handwritten correspondence requiring conservation treatment before digitization."
    )
    doc.add_heading("Common Themes Across Collections", level=2)
    doc.add_paragraph(
        "Key themes: political transitions, diplomatic correspondence, social movements, "
        "economic development, cultural heritage documentation. "
        "Cross-cutting theme: the role of international institutions in MENA governance."
    )
    doc.save(summary_dir / "collection_assessment.docx")

    print(f"Created 3 Word documents in {summary_dir}")

except ImportError:
    print("python-docx not installed — skipping Word test data")

print("\nTest data creation complete.")
print(f"  Financial data: {yoy_dir}")
print(f"  Report data: {summary_dir}")
